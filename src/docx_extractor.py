#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
import os
from typing import List, Tuple, Dict, Set
import logging
import re
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import threading
import sys

# 版本信息
try:
    from src import __version__, __app_name__
except ImportError:
    __version__ = "2.1.0"
    __app_name__ = "DOCX文档提取器"

# 字符宽度相关常量
UNICODE_BOUNDARY = 127  # ASCII和Unicode字符的分界点
SINGLE_CHAR_WIDTH = 1   # ASCII字符宽度
DOUBLE_CHAR_WIDTH = 2   # Unicode字符（如中文）宽度

# 文本格式化相关常量
DEFAULT_MAX_WIDTH = 80      # 默认每行最大字符数
DEFAULT_INDENT = "    "     # 默认缩进（4个空格）
HEADING_PREFIX = "#"        # 标题前缀符号

# 表格相关常量
# 列宽度控制
MIN_COLUMN_WIDTH = 8        # 表格列最小宽度（字符数）
NORMAL_MAX_WIDTH = 20       # 普通列的最大宽度
MULTILINE_MAX_WIDTH = 40    # 多行文本列的最大宽度
CELL_PADDING = 2           # 单元格内容两侧的空白padding总和
LONG_TEXT_THRESHOLD = 25    # 判定为长文本的宽度阈值（字符数）
CELL_LEFT_PADDING = 1      # 单元格左侧padding（字符数）
NEWLINE_THRESHOLD = 2      # 触发列宽加倍的换行次数阈值

# 添加新的常量定义
BASE_COLUMN_WIDTH = 15  # 基础列宽
LEVEL_2_MULTIPLIER = 2  # Level 2 宽度倍数
LEVEL_3_MULTIPLIER = 3  # Level 3 宽度倍数

class DocxExtractor:
    """DOCX文本提取器类，用于提取Word文档中的文本和表格内容，保持原文档的结构和位置关系"""
    
    def __init__(self, docx_path: str, auto_setup_logging: bool = True):
        """
        初始化DOCX提取器
        
        Args:
            docx_path: DOCX文件路径
            auto_setup_logging: 是否自动设置日志系统
        """
        self.docx_path = docx_path
        self.output_path = os.path.splitext(docx_path)[0] + '.md'
        if auto_setup_logging:
            self._setup_logging()
        
    def _setup_logging(self, gui_handler=None):
        """配置日志系统，仅输出到控制台"""
        # 创建简洁的格式化器
        formatter = logging.Formatter('%(levelname)s: %(message)s')
        
        # 创建控制台处理器
        console_handler = logging.StreamHandler()
        console_handler.setFormatter(formatter)
        
        # 获取logger实例
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(logging.INFO)
        
        # 清除已有的处理器
        self.logger.handlers.clear()
        
        # 添加控制台处理器
        self.logger.addHandler(console_handler)
        
        # 如果提供了GUI处理器，也添加它
        if gui_handler:
            self.logger.addHandler(gui_handler)
    
    def _get_char_width(self, char: str) -> int:
        """
        获取字符的显示宽度
        
        Args:
            char: 单个字符
        
        Returns:
            int: 字符的显示宽度
        """
        return DOUBLE_CHAR_WIDTH if ord(char) > UNICODE_BOUNDARY else SINGLE_CHAR_WIDTH
    
    def _get_string_width(self, text: str) -> int:
        """
        获取字符串的显示宽度
        
        Args:
            text: 字符串
        
        Returns:
            int: 字符串的显示宽度
        """
        return sum(self._get_char_width(char) for char in str(text))
    
    def _clean_text(self, text: str) -> str:
        """
        清理文本内容，保留必要的格式，合并连续空行为一行
        
        Args:
            text: 原始文本
            
        Returns:
            str: 清理后的文本
        """
        if not text:
            return ""
        
        # 保留换行符，但删除每行开头和结尾的空白字符
        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines]
        
        # 合并连续的空行为单个空行
        result_lines = []
        prev_empty = False
        
        for line in cleaned_lines:
            is_empty = not line
            
            if is_empty:
                # 只有当前一行不是空行时，才添加空行
                if not prev_empty:
                    result_lines.append(line)
                prev_empty = True
            else:
                # 非空行直接添加
                result_lines.append(line)
                prev_empty = False
        
        return '\n'.join(result_lines)
    
    def _process_heading(self, paragraph) -> str:
        """
        处理标题段落
        
        Args:
            paragraph: 段落对象
            
        Returns:
            str: 格式化的标题文本
        """
        self.logger.debug(f"开始处理标题段落: {paragraph.text[:50]}...")
        
        text = paragraph.text.strip()
        if not text:
            self.logger.warning("标题段落为空")
            return ""
            
        # 根据标题级别添加不同数量的#
        level = paragraph.style.name.lower()
        if 'heading' in level:
            try:
                level_num = int(level[-1])
                result = f"{'#' * level_num} {text}\n"
                self.logger.debug(f"处理标题完成，级别: {level_num}, 内容: {text[:50]}...")
                return result
            except ValueError:
                self.logger.warning(f"无法解析标题级别: {level}")
                
        result = f"# {text}\n"
        self.logger.debug(f"处理默认标题完成: {text[:50]}...")
        return result
    
    def _wrap_text_by_width(self, text: str, max_width: int = DEFAULT_MAX_WIDTH, indent: str = "") -> str:
        """
        按照视觉宽度对文本自动换行
        
        Args:
            text: 要处理的文本
            max_width: 每行最大宽度（按英文字符计算）
            indent: 缩进字符串
            
        Returns:
            str: 处理后的文本
        """
        if not text:
            return ""
            
        # 计算缩进占用的宽度
        indent_width = self._get_string_width(indent)
        # 实际可用宽度
        available_width = max_width - indent_width
        
        result_lines = []
        current_line = []
        current_width = 0
        
        # 分词处理，保持英文单词的完整性
        words = []
        temp_word = ""
        
        for char in text:
            if char.isspace():
                if temp_word:
                    words.append(temp_word)
                    temp_word = ""
                words.append(char)
            else:
                # 对于中文字符，单独作为一个词
                if ord(char) > 127:
                    if temp_word:
                        words.append(temp_word)
                        temp_word = ""
                    words.append(char)
                else:
                    temp_word += char
        if temp_word:
            words.append(temp_word)
            
        # 按宽度分行
        for word in words:
            word_width = self._get_string_width(word)
            
            # 如果当前行是空的，直接添加
            if not current_line:
                current_line.append(word)
                current_width = word_width
                continue
                
            # 检查添加这个词是否会超出宽度
            if current_width + word_width <= available_width:
                current_line.append(word)
                current_width += word_width
            else:
                # 当前行已满，开始新行
                result_lines.append("".join(current_line))
                current_line = [word]
                current_width = word_width
                
        # 处理最后一行
        if current_line:
            result_lines.append("".join(current_line))
            
        # 添加缩进并连接所有行
        return "\n".join(indent + line.strip() for line in result_lines)
        
    def _process_normal_paragraph(self, paragraph) -> str:
        """
        处理普通段落，实现基于视觉宽度的自动换行
        
        Args:
            paragraph: 段落对象
            
        Returns:
            str: 格式化的段落文本
        """
        self.logger.debug(f"开始处理普通段落: {paragraph.text[:50]}...")
        
        if not paragraph.text:
            self.logger.debug("段落为空，返回换行符")
            return "\n"
            
        text = paragraph.text.strip()
        
        if paragraph.style and 'heading' in paragraph.style.name.lower():
            try:
                level = int(paragraph.style.name.lower()[-1])
                prefix = HEADING_PREFIX * level + ' '
                result = f"{prefix}{text}\n\n"
                self.logger.debug(f"处理带样式段落完成，样式: {paragraph.style.name}")
                return result
            except ValueError:
                prefix = HEADING_PREFIX + ' '
                self.logger.warning(f"无法解析段落样式级别: {paragraph.style.name}")
                result = f"{prefix}{text}\n\n"
                return result
            
        wrapped_text = self._wrap_text_by_width(text, max_width=100, indent=DEFAULT_INDENT)
        self.logger.debug(f"处理普通段落完成，长度: {len(text)}")
        return wrapped_text + "\n\n"
    
    def _get_merged_cell_info(self, cell, row_idx: int, col_idx: int) -> Tuple[int, int]:
        """
        获取单元格的合并信息，通过多个条件判断是否应该合并
        
        Args:
            cell: 单元格对象
            row_idx: 行索引
            col_idx: 列索引
            
        Returns:
            Tuple[int, int]: (垂直合并跨度, 水平合并跨度)
        """
        vmerge = cell._tc.get_or_add_tcPr().first_child_found_in("w:vMerge")
        hmerge = cell._tc.get_or_add_tcPr().first_child_found_in("w:gridSpan")
        
        vspan = 1
        hspan = int(hmerge.val) if hmerge is not None else 1
        
        # 处理垂直合并
        if vmerge is not None:
            current_text = cell.text.strip()
            
            if vmerge.val == "restart":  # 合并的起始单元格
                # 找出合并的行数
                current_row = row_idx + 1
                while current_row < len(self.current_table.rows):
                    next_cell = self.current_table.rows[current_row].cells[col_idx]
                    next_vmerge = next_cell._tc.get_or_add_tcPr().first_child_found_in("w:vMerge")
                    next_text = next_cell.text.strip()
                    
                    # 判断是否应该合并的条件：
                    # 1. 下一个单元格有vMerge属性
                    # 2. 不是新的合并起始点
                    # 3. 如果当前单元格有内容，则要求下一个单元格内容相同或为空
                    # 4. 如果当前单元格为空，则不要求下一个单元格内容
                    should_merge = (next_vmerge is not None and 
                                  (next_vmerge.val != "restart") and
                                  (not current_text or 
                                   not next_text or 
                                   current_text == next_text))
                    
                    if not should_merge:
                        break
                        
                    vspan += 1
                    current_row += 1
                    
            else:  # 被合并的单元格
                # 向上查找合并的起始单元格
                current_row = row_idx - 1
                found_start = False
                
                while current_row >= 0:
                    prev_cell = self.current_table.rows[current_row].cells[col_idx]
                    prev_vmerge = prev_cell._tc.get_or_add_tcPr().first_child_found_in("w:vMerge")
                    prev_text = prev_cell.text.strip()
                    
                    # 判断是否是合并的起始点：
                    # 1. 上一个单元格有vMerge属性且是restart
                    # 2. 如果当前单元格有内容，则要求与起始单元格内容相同
                    # 3. 如果当前单元格为空，则不要求起始单元格内容
                    if (prev_vmerge is not None and 
                        prev_vmerge.val == "restart" and
                        (not current_text or 
                         not prev_text or 
                         current_text == prev_text)):
                        found_start = True
                        vspan = 0  # 标记为被合并的单元格
                        break
                        
                    current_row -= 1
                    
                # 如果向上没有找到合适的起始点，则不进行合并
                if not found_start:
                    vspan = 1
                
        return vspan, hspan
        
    def _process_cell_content(self, cell_text: str, total_width: int) -> List[str]:
        """
        处理单元格内容，进行文本换行
        
        Args:
            cell_text: 单元格文本
            total_width: 单元格总宽度
            
        Returns:
            List[str]: 处理后的文本行列表
        """
        if not cell_text:
            return ['']
        
        wrapped_lines = []
        available_width = total_width - CELL_PADDING
        
        # 分行处理
        lines = cell_text.split('\n')
        
        for line in lines:
            if not line:
                wrapped_lines.append('')
                continue
            
            current_line = []
            current_width = 0
            
            # 分词处理
            words = []
            temp_word = ""
            
            for char in line:
                if char.isspace():
                    if temp_word:
                        words.append(temp_word)
                        temp_word = ""
                    words.append(char)
                else:
                    if ord(char) > 127:  # 中文字符
                        if temp_word:
                            words.append(temp_word)
                            temp_word = ""
                        words.append(char)
                    else:
                        temp_word += char
            if temp_word:
                words.append(temp_word)
            
            # 组织行
            for word in words:
                word_width = self._get_string_width(word)
                
                if current_width + word_width <= available_width:
                    current_line.append(word)
                    current_width += word_width
                else:
                    if current_line:
                        wrapped_lines.append(''.join(current_line))
                    current_line = [word]
                    current_width = word_width
            
            if current_line:
                wrapped_lines.append(''.join(current_line))
        
        return wrapped_lines if wrapped_lines else ['']
        
    def _calculate_column_widths(self, table_data: List[List[Dict]], max_cols: int) -> List[int]:
        """
        计算表格每列的宽度
        
        Args:
            table_data: 表格数据
            max_cols: 最大列数
            
        Returns:
            List[int]: 每列的宽度列表
        """
        self.logger.debug("开始计算列宽...")
        
        # 初始化列宽数组
        col_widths = [0] * max_cols
        
        # 第一步:计算每列中最大内容宽度
        max_content_widths = [0] * max_cols
        for row_data in table_data:
            for col_idx, cell_data in enumerate(row_data):
                if cell_data['hspan'] == 1:  # 只处理非合并单元格
                    content_width = cell_data['max_line_width']
                    max_content_widths[col_idx] = max(max_content_widths[col_idx], content_width)
        
        # 第二步:根据内宽度确定每列等级
        for col_idx in range(max_cols):
            max_width = max_content_widths[col_idx]
            if max_width <= BASE_COLUMN_WIDTH:
                col_widths[col_idx] = BASE_COLUMN_WIDTH
            elif max_width <= BASE_COLUMN_WIDTH * LEVEL_2_MULTIPLIER:
                col_widths[col_idx] = BASE_COLUMN_WIDTH * LEVEL_2_MULTIPLIER
            elif max_width <= BASE_COLUMN_WIDTH * LEVEL_3_MULTIPLIER:
                col_widths[col_idx] = BASE_COLUMN_WIDTH * LEVEL_3_MULTIPLIER
            else:
                col_widths[col_idx] = BASE_COLUMN_WIDTH * LEVEL_3_MULTIPLIER
            
        # 第三步:处理合并单元格
        for row_data in table_data:
            for col_idx, cell_data in enumerate(row_data):
                if cell_data['hspan'] > 1:
                    # 计算合并单元格占用的总宽度
                    total_width = sum(col_widths[i] for i in range(col_idx, min(col_idx + cell_data['hspan'], len(col_widths))))
                    # 更新cell_data中的实际可用宽度信息
                    cell_data['available_width'] = total_width - CELL_PADDING
        
        self.logger.debug(f"列宽计算完成: {col_widths}")
        return col_widths
        
    def _format_table_row(self, row_data: List[Dict], col_widths: List[int], line_idx: int) -> str:
        """
        格式化表格行，支持垂直合并单元格
        
        Args:
            row_data: 行数据
            col_widths: 列宽列表
            line_idx: 行索引
            
        Returns:
            str: 格式化后的行字符串
        """
        row_content = '|'
        col_idx = 0
        
        while col_idx < len(row_data):
            cell_data = row_data[col_idx]
            hspan = cell_data['hspan']
            
            # 计算单元格总宽度
            total_width = sum(col_widths[i] for i in range(col_idx, min(col_idx + hspan, len(col_widths)))) + (hspan - 1)
            
            # 获取单元格文本
            if cell_data['vspan'] == 0:  # 被合并的单元格
                cell_text = ''
            else:
                cell_text = cell_data['lines'][line_idx] if line_idx < len(cell_data['lines']) else ''
            
            # 计算填充
            content_width = self._get_string_width(cell_text)
            left_padding = CELL_LEFT_PADDING
            right_padding = total_width - content_width - left_padding
            
            # 添加单元格内容
            row_content += ' ' * left_padding + cell_text + ' ' * right_padding + '|'
            
            col_idx += hspan
        
        return row_content
        
    def _collect_table_data(self, table) -> Tuple[List[List[Dict]], int]:
        """
        收集表格数据和合并单元格信息
        
        Args:
            table: 表格对象
            
        Returns:
            Tuple[List[List[Dict]], int]: (表格数据, 最大列数)
        """
        self.logger.debug("开始收集表格数据...")
        max_cols = max(len(row.cells) for row in table.rows)
        table_data = []
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for col_idx, cell in enumerate(row.cells):
                vspan, hspan = self._get_merged_cell_info(cell, row_idx, col_idx)
                
                if cell._tc.get_or_add_tcPr().first_child_found_in("w:vMerge") is not None and \
                   cell._tc.get_or_add_tcPr().first_child_found_in("w:vMerge").val != "restart":
                    continue
                
                # 清理并保持单元格内的换行
                cell_text = self._clean_text(cell.text)
                original_lines = cell_text.split('\n')
                line_widths = [self._get_string_width(line) for line in original_lines]
                max_line_width = max(line_widths) if line_widths else 0
                
                # 判断是否需要加宽列
                needs_double_width = len(original_lines) - 1 >= NEWLINE_THRESHOLD
                
                row_data.append({
                    'text': cell_text,
                    'lines': original_lines,
                    'hspan': hspan,
                    'vspan': vspan,
                    'needs_double_width': needs_double_width,
                    'max_line_width': max_line_width
                })
            
            # 补齐缺失的列
            while len(row_data) < max_cols:
                row_data.append({
                    'text': '',
                    'lines': [''],
                    'hspan': 1,
                    'vspan': 1,
                    'needs_double_width': False,
                    'max_line_width': 0
                })
                
            table_data.append(row_data)
            
        self.logger.debug(f"表格数据收集完成，共 {len(table_data)} 行")
        return table_data, max_cols

    def _process_cell_wrapping(self, table_data: List[List[Dict]], col_widths: List[int]):
        """
        处理所有单元格的文本换行，包括纵向合并单元格
        
        Args:
            table_data: 表格数据
            col_widths: 列宽列表
        """
        self.logger.debug("开始处理单元格文本换行...")
        
        # 第一遍：处理垂直合并单元格的文本
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                if cell_data['vspan'] > 1:  # 处理垂直合并单元格
                    # 计算水平方向的总宽度
                    h_total_width = sum(col_widths[i] for i in range(col_idx, min(col_idx + cell_data['hspan'], len(col_widths)))) + (cell_data['hspan'] - 1)
                    
                    # 处理合并单元格的文本换行
                    if cell_data['text']:
                        cell_data['lines'] = self._process_cell_content(cell_data['text'], h_total_width)
                        
                    # 标记被合并的单元格
                    for v_idx in range(row_idx + 1, min(row_idx + cell_data['vspan'], len(table_data))):
                        table_data[v_idx][col_idx] = {
                            'text': '',
                            'lines': [''],  # 使用空字符串而不是空列表
                            'hspan': cell_data['hspan'],
                            'vspan': 0,  # 标记为被合并的单元格
                            'needs_double_width': False,
                            'max_line_width': 0,
                            'merged_from': (row_idx, col_idx)  # 添加合并源信息
                        }
        
        # 第二遍：处理普通单元格
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                if cell_data['vspan'] == 1:  # 只处理非垂直合并单元格
                    if not cell_data['text']:
                        cell_data['lines'] = ['']  # 确保空单元格也有一个空行
                        continue
                    
                    total_width = sum(col_widths[i] for i in range(col_idx, min(col_idx + cell_data['hspan'], len(col_widths)))) + (cell_data['hspan'] - 1)
                    cell_data['lines'] = self._process_cell_content(cell_data['text'], total_width)
        
        self.logger.debug("单元格文本换行处理完成")

    def _generate_table_string(self, table_data: List[List[Dict]], col_widths: List[int]) -> str:
        """
        生成最终的表格字符串，支持垂直合并单元格
        
        Args:
            table_data: 表格数据
            col_widths: 列宽列表
            
        Returns:
            str: 格式化的表格字符串
        """
        self.logger.debug("开始生成表格字符串...")
        result = []
        
        # 生成分隔线
        separator = '+'
        for width in col_widths:
            separator += '-' * width + '+'
        
        # 处理每一行
        for row_idx, row_data in enumerate(table_data):
            # 获取当前行中所有非被合并单元格的最大行数
            max_lines = 1  # 至少有一行
            for cell_data in row_data:
                if cell_data['vspan'] != 0:  # 不是被合并的单元格
                    max_lines = max(max_lines, len(cell_data['lines']))
            
            # 补齐每个单元格的行数
            for cell_data in row_data:
                if cell_data['vspan'] != 0:  # 不是被合并的单元格
                    # 确保lines至少有一个元素
                    if not cell_data['lines']:
                        cell_data['lines'] = ['']
                    # 补齐行数
                    while len(cell_data['lines']) < max_lines:
                        cell_data['lines'].append('')
            
            # 添加分隔线
            # 只在以下情况添加分隔线：
            # 1. 第一行
            # 2. 当前行没有被合并的单元格
            # 3. 当前行是某个垂直合并单元格的起始行
            if (row_idx == 0 or 
                not any(cell_data['vspan'] == 0 for cell_data in row_data) or
                any(cell_data['vspan'] > 1 for cell_data in row_data)):
                result.append(separator)
            
            # 处理多行单元格
            for line_idx in range(max_lines):
                result.append(self._format_table_row(row_data, col_widths, line_idx))
        
        # 添加最后一行分隔线
        result.append(separator)
        
        table_string = '\n'.join(result) + '\n'
        self.logger.debug(f"表格字符串生成完成总行数: {len(result)}")
        return table_string

    def _process_table(self, table) -> str:
        """
        处理表格，保持原始格式和对齐方式，支持合并单元格和自动换行
        
        Args:
            table: 表格对象
            
        Returns:
            str: 格式化的表格文本
        """
        try:
            self.logger.info(f"处理表格: {len(table.rows)}行 x {len(table.rows[0].cells) if table.rows else 0}列")
            
            if not table.rows:
                self.logger.warning("表格为空")
                return "\n"
                
            self.current_table = table  # 存储当前处理的表格对象
            
            # 1. 收集表格数据
            table_data, max_cols = self._collect_table_data(table)
            
            # 2. 计算列宽
            col_widths = self._calculate_column_widths(table_data, max_cols)
            
            # 3. 处理文本换行
            self._process_cell_wrapping(table_data, col_widths)
            
            # 4. 生成表格字符串
            result = self._generate_table_string(table_data, col_widths)
            self.logger.info("表格处理完成")
            
            return result
            
        except Exception as e:
            self.logger.error(f"处理表格时出错: {str(e)}", exc_info=True)
            return "【表格处理失败】\n"
    
    def extract_and_save(self):
        """提取DOCX内容并保存到文本文件，保持原文档结构"""
        try:
            self.logger.info(f"开始处理DOCX文件: {self.docx_path}")
            
            doc = Document(self.docx_path)
            content = []
            
            # 获取所有段落和表格
            paragraphs = list(doc.paragraphs)
            tables = list(doc.tables)
            p_index = 0
            t_index = 0
            
            # 理文档中的每个元素
            total_elements = len(list(doc.element.body))
            processed_elements = 0
            
            for element in doc.element.body:
                if element.tag.endswith('p'):  # 段落
                    if p_index < len(paragraphs):
                        paragraph = paragraphs[p_index]
                        if paragraph.style and paragraph.style.name and 'heading' in paragraph.style.name.lower():
                            content.append(self._process_heading(paragraph))
                        else:
                            content.append(self._process_normal_paragraph(paragraph))
                        p_index += 1
                elif element.tag.endswith('tbl'):  # 表格
                    if t_index < len(tables):
                        content.append(self._process_table(tables[t_index]))
                        t_index += 1
                
                processed_elements += 1
                if processed_elements % 10 == 0 or processed_elements == total_elements:
                    self.logger.info(f"处理进度: {processed_elements}/{total_elements} ({processed_elements/total_elements*100:.1f}%)")
                
            # 保存到文本文件
            with open(self.output_path, 'w', encoding='utf-8') as f:
                f.write(''.join(content))
            
            self.logger.info(f"文件处理完成，已保存到: {self.output_path}")
            return True
            
        except FileNotFoundError:
            self.logger.error(f"找不到DOCX文件: {self.docx_path}")
            raise
        except KeyError as e:
            # 处理docx内部损坏的引用（如断开的书签链接）
            error_msg = str(e)
            if 'word/' in error_msg and 'bookmark' in error_msg.lower():
                self.logger.warning(f"文件包含损坏的书签引用，已跳过: {error_msg}")
            else:
                self.logger.warning(f"文件包含损坏的内部引用，已跳过: {error_msg}")
            # 尝试保存已处理的内容
            if content:
                self.logger.info(f"已提取部分内容，正在保存...")
                with open(self.output_path, 'w', encoding='utf-8') as f:
                    f.write(''.join(content))
                self.logger.info(f"部分内容已保存到: {self.output_path}")
                return True
            raise
        except Exception as e:
            self.logger.error(f"处理DOCX时发生错误: {str(e)}")
            raise

class GUILogHandler(logging.Handler):
    """GUI日志处理器，将日志输出到文本框，带缓冲机制"""
    
    def __init__(self, text_widget):
        super().__init__()
        self.text_widget = text_widget
        self.buffer = []
        self.update_pending = False
        
    def emit(self, record):
        """输出日志记录到文本框"""
        try:
            msg = self.format(record)
            self.buffer.append(msg + '\n')
            
            # 如果没有待处理的更新，安排一个
            if not self.update_pending:
                self.update_pending = True
                self.text_widget.after(100, self._flush_buffer)
        except Exception:
            pass
            
    def _flush_buffer(self):
        """批量更新文本框"""
        try:
            if self.buffer:
                # 批量插入所有缓冲的日志
                text = ''.join(self.buffer)
                self.text_widget.insert(tk.END, text)
                self.text_widget.see(tk.END)
                self.buffer.clear()
            self.update_pending = False
        except Exception:
            self.update_pending = False

class DocxExtractorGUI:
    """DOCX提取器图形界面"""
    
    def __init__(self):
        self.root = tk.Tk()
        self.root.title(f"{__app_name__} v{__version__}")
        self.root.geometry("800x600")
        self.root.resizable(True, True)
        
        # 设置窗口图标
        self._set_window_icon()
        
        self.setup_ui()
    
    def _set_window_icon(self):
        """设置窗口图标"""
        try:
            if getattr(sys, 'frozen', False):
                base_path = sys._MEIPASS
            else:
                base_path = os.path.dirname(os.path.abspath(__file__))
            
            icon_path = os.path.join(base_path, 'assets', 'app_icon.ico')
            if os.path.exists(icon_path):
                self.root.iconbitmap(icon_path)
        except Exception:
            pass
        
    def setup_ui(self):
        """设置用户界面"""
        # 主框架
        main_frame = tk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 文件选择框架
        file_frame = tk.Frame(main_frame)
        file_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 文件路径标签和输入框
        tk.Label(file_frame, text="选择DOCX文件:").pack(side=tk.LEFT)
        self.file_path_var = tk.StringVar()
        self.file_entry = tk.Entry(file_frame, textvariable=self.file_path_var, width=50)
        self.file_entry.pack(side=tk.LEFT, padx=(10, 5), fill=tk.X, expand=True)
        
        # 浏览按钮
        browse_btn = tk.Button(file_frame, text="浏览", command=self.browse_file)
        browse_btn.pack(side=tk.LEFT, padx=(5, 10))
        
        # 处理按钮
        process_btn = tk.Button(file_frame, text="开始处理", command=self.process_file)
        process_btn.pack(side=tk.RIGHT)
        
        # 日志显示区域
        log_frame = tk.Frame(main_frame)
        log_frame.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(log_frame, text="处理日志:").pack(anchor=tk.W)
        
        # 创建带滚动条的文本框
        self.log_text = scrolledtext.ScrolledText(
            log_frame, 
            wrap=tk.WORD, 
            width=80, 
            height=25,
            font=("Consolas", 9)
        )
        self.log_text.pack(fill=tk.BOTH, expand=True, pady=(5, 0))
        
        # 状态栏
        self.status_var = tk.StringVar()
        self.status_var.set("就绪")
        status_bar = tk.Label(
            self.root, 
            textvariable=self.status_var, 
            relief=tk.SUNKEN, 
            anchor=tk.W
        )
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
    def browse_file(self):
        """浏览文件对话框"""
        file_path = filedialog.askopenfilename(
            title="选择DOCX文件",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if file_path:
            self.file_path_var.set(file_path)
            
    def process_file(self):
        """处理文件"""
        file_path = self.file_path_var.get().strip()
        
        if not file_path:
            messagebox.showerror("错误", "请选择一个DOCX文件")
            return
            
        if not os.path.exists(file_path):
            messagebox.showerror("错误", "文件不存在")
            return
            
        if not file_path.lower().endswith('.docx'):
            messagebox.showerror("错误", "请选择DOCX格式的文件")
            return
            
        # 清空日志区域
        self.log_text.delete(1.0, tk.END)
        
        # 在新线程中处理文件，避免阻塞GUI
        threading.Thread(target=self._process_file_thread, args=(file_path,), daemon=True).start()
        
    def _process_file_thread(self, file_path):
        """在后台线程中处理文件"""
        try:
            self.status_var.set("处理中...")
            
                         # 创建提取器实例，不自动设置日志
            extractor = DocxExtractor(file_path, auto_setup_logging=False)
            
            # 创建GUI日志处理器
            gui_handler = GUILogHandler(self.log_text)
            gui_handler.setFormatter(logging.Formatter('%(levelname)s: %(message)s'))
            gui_handler.setLevel(logging.INFO)  # 只显示INFO及以上级别的日志
            
            # 设置日志系统，包含GUI处理器
            extractor._setup_logging(gui_handler)
            
            # 开始处理
            extractor.extract_and_save()
            
            # 处理完成
            self.status_var.set("处理完成")
            
            # 显示成功消息
            output_file = os.path.splitext(file_path)[0] + '.md'
            self.root.after(0, lambda: messagebox.showinfo(
                "成功", 
                f"文件处理完成!\n输出文件: {output_file}"
            ))
            
        except Exception as e:
            self.status_var.set("处理失败")
            error_msg = f"处理文件时发生错误: {str(e)}"
            self.log_text.after(0, self._append_error, error_msg)
            self.root.after(0, lambda: messagebox.showerror("错误", error_msg))
            
    def _append_error(self, error_msg):
        """在日志中追加错误信息"""
        self.log_text.insert(tk.END, f"ERROR: {error_msg}\n")
        self.log_text.see(tk.END)
        
    def run(self):
        """运行GUI"""
        self.root.mainloop()

def main():
    """主函数"""
    # 检查是否有命令行参数
    if len(sys.argv) == 2:
        # 命令行模式
        docx_path = sys.argv[1]
        extractor = DocxExtractor(docx_path)
        
        try:
            extractor.extract_and_save()
        except Exception as e:
            print(f"错误: {str(e)}")
            sys.exit(1)
    else:
        # GUI模式
        try:
            app = DocxExtractorGUI()
            app.run()
        except Exception as e:
            print(f"启动GUI失败: {str(e)}")
            print("使用方法: python docx_extractor.py <docx文件路径>")
            sys.exit(1)

if __name__ == "__main__":
    main() 