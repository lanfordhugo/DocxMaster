#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
DOCX文档提取器核心模块
负责文档解析和文本提取的核心逻辑
"""

import logging
import os
import re
import hashlib
from typing import Dict, List, Tuple, Optional, Any

from docx import Document

# 字符宽度相关常量
UNICODE_BOUNDARY = 127  # ASCII和Unicode字符的分界点
SINGLE_CHAR_WIDTH = 1   # ASCII字符宽度
DOUBLE_CHAR_WIDTH = 2   # Unicode字符（如中文）宽度

# 文本格式化相关常量
DEFAULT_MAX_WIDTH = 80      # 默认每行最大字符数
DEFAULT_INDENT = "    "     # 默认缩进（4个空格）
HEADING_PREFIX = "#"        # 标题前缀符号

# 表格相关常量
MIN_COLUMN_WIDTH = 8        # 表格列最小宽度（字符数）
NORMAL_MAX_WIDTH = 20       # 普通列的最大宽度
MULTILINE_MAX_WIDTH = 40    # 多行文本列的最大宽度
CELL_PADDING = 2           # 单元格内容两侧的空白padding总和
LONG_TEXT_THRESHOLD = 25    # 判定为长文本的宽度阈值（字符数）
CELL_LEFT_PADDING = 1      # 单元格左侧padding（字符数）
NEWLINE_THRESHOLD = 2      # 触发列宽加倍的换行次数阈值

# 列宽度控制常量
BASE_COLUMN_WIDTH = 15  # 基础列宽
LEVEL_2_MULTIPLIER = 2  # Level 2 宽度倍数
LEVEL_3_MULTIPLIER = 3  # Level 3 宽度倍数


class DocumentExtractor:
    """DOCX文档提取器核心类"""
    
    def __init__(self, config: Optional[Any] = None) -> None:
        """初始化提取器
        
        Args:
            config: 可选配置对象，需包含 text_width/text_indent 等属性
        """
        self.logger: logging.Logger = logging.getLogger(__name__)
        self.current_table = None
        # 配置相关（带默认值，避免强耦合），固定输出md格式
        self.text_width: int = getattr(config, 'text_width', 100) if config else 100
        self.text_indent: str = getattr(config, 'text_indent', DEFAULT_INDENT) if config else DEFAULT_INDENT
        
    def _get_char_width(self, char: str) -> int:
        """获取字符的显示宽度"""
        return DOUBLE_CHAR_WIDTH if ord(char) > UNICODE_BOUNDARY else SINGLE_CHAR_WIDTH
    
    def _get_string_width(self, text: str) -> int:
        """获取字符串的显示宽度
        
        Args:
            text: 输入字符串
            
        Returns:
            字符串的显示宽度
        """
        return sum(self._get_char_width(char) for char in str(text))
    
    def _clean_text(self, text: str) -> str:
        """清理文本内容，保留必要的格式，合并连续空行为一行"""
        if not text:
            return ""
        
        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines]
        
        # 合并连续的空行为单个空行
        result_lines = []
        prev_empty = False
        
        for line in cleaned_lines:
            is_empty = not line
            
            if is_empty:
                if not prev_empty:
                    result_lines.append(line)
                prev_empty = True
            else:
                result_lines.append(line)
                prev_empty = False
        
        return '\n'.join(result_lines)
    
    def _process_heading(self, paragraph) -> str:
        """处理标题段落（在 md 模式下注入稳定锚点）"""
        text = paragraph.text.strip()
        if not text:
            return ""
        
        # 标题级别
        level = paragraph.style.name.lower()
        level_num = None
        if 'heading' in level:
            try:
                level_num = int(level[-1])
            except ValueError:
                level_num = None

        prefix = (HEADING_PREFIX * level_num + ' ') if level_num else (HEADING_PREFIX + ' ')

        # 规范 cmd 标注与锚点
        norm_text = text.replace('（', '(').replace('）', ')')
        # 归一化 cmd 大小写与空格
        norm_text = re.sub(r'\bCMD\b', 'cmd', norm_text, flags=re.I)
        cmd_match = re.search(r'cmd\s*=\s*(\d+)', norm_text, flags=re.I)
        anchor_line = ''
        if cmd_match:
            cmd_val = int(cmd_match.group(1))
            anchor_id = f"cmd-{cmd_val:03d}"
            # 若文本未带标准 [cmd=xxx]，追加标准化标注
            if not re.search(r'\[\s*cmd\s*=\s*\d+\s*\]', norm_text, flags=re.I):
                norm_text = f"{norm_text} [cmd={cmd_val:03d}]"
            anchor_line = f"<a id=\"{anchor_id}\"></a>\n"
        else:
            # 为无 cmd 的标题生成稳定锚点（基于内容的短哈希）
            digest = hashlib.sha1(norm_text.encode('utf-8')).hexdigest()[:8]
            anchor_id = f"sec-{digest}"
            anchor_line = f"<a id=\"{anchor_id}\"></a>\n"
        return f"{anchor_line}{prefix}{norm_text}\n"
    
    def _process_pseudo_cmd_title(self, text: str) -> str:
        """处理伪CMD标题（普通段落中识别的CMD格式）"""
        # 规范化文本：全角→半角，统一cmd大小写
        norm_text = text.replace('（', '(').replace('）', ')')
        norm_text = re.sub(r'\bCMD\b', 'cmd', norm_text, flags=re.I)
        
        # 提取 CMD 编号
        cmd_match = re.search(r'cmd\s*=\s*(\d+)', norm_text, flags=re.I)
        anchor_line = ''
        if cmd_match:
            cmd_val = int(cmd_match.group(1))
            anchor_id = f"cmd-{cmd_val:03d}"
            # 若文本未带标准 [cmd=xxx]，追加标准化标注
            if not re.search(r'\[\s*cmd\s*=\s*\d+\s*\]', norm_text, flags=re.I):
                norm_text = f"{norm_text} [cmd={cmd_val:03d}]"
            anchor_line = f"<a id=\"{anchor_id}\"></a>\n"
        else:
            # 生成基于内容的哈希锚点（备用）
            digest = hashlib.sha1(norm_text.encode('utf-8')).hexdigest()[:8]
            anchor_id = f"sec-{digest}"
            anchor_line = f"<a id=\"{anchor_id}\"></a>\n"
        
        # 作为三级标题输出
        return f"{anchor_line}### {norm_text}\n\n"
    
    def _wrap_text_by_width(self, text: str, max_width: int = DEFAULT_MAX_WIDTH, indent: str = "") -> str:
        """按照视觉宽度对文本自动换行"""
        if not text:
            return ""
            
        indent_width = self._get_string_width(indent)
        available_width = max_width - indent_width
        
        result_lines = []
        current_line = []
        current_width = 0
        
        # 分词处理
        words = []
        temp_word = ""
        
        for char in text:
            if char.isspace():
                if temp_word:
                    words.append(temp_word)
                    temp_word = ""
                words.append(char)
            else:
                if ord(char) > 127:  # 中文字符单独作为词
                    if temp_word:
                        words.append(temp_word)
                        temp_word = ""
                    words.append(char)
                else:
                    temp_word += char
        if temp_word:
            words.append(temp_word)
            
        # 按宽度分行
        for word in words:
            word_width = self._get_string_width(word)
            
            if not current_line:
                current_line.append(word)
                current_width = word_width
                continue
                
            if current_width + word_width <= available_width:
                current_line.append(word)
                current_width += word_width
            else:
                result_lines.append("".join(current_line))
                current_line = [word]
                current_width = word_width
                
        if current_line:
            result_lines.append("".join(current_line))
            
        return "\n".join(indent + line.strip() for line in result_lines)
    
    def _process_normal_paragraph(self, paragraph) -> str:
        """处理普通段落"""
        if not paragraph.text:
            return "\n"
            
        text = paragraph.text.strip()
        
        if paragraph.style and 'heading' in paragraph.style.name.lower():
            try:
                level = int(paragraph.style.name.lower()[-1])
                prefix = HEADING_PREFIX * level + ' '
                # 走标题处理，保证锚点一致
                return self._process_heading(paragraph) + "\n"
            except ValueError:
                prefix = HEADING_PREFIX + ' '
                return f"{prefix}{text}\n\n"
        
        # 识别形如 "x.x.x (CMD=xxx)" 的伪标题段落
        cmd_pattern = r'^\s*\d+\.\d+(?:\.\d+)?\s*\([Cc][Mm][Dd]\s*=\s*\d+\)'
        if re.match(cmd_pattern, text):
            # 将伪标题按三级标题处理
            return self._process_pseudo_cmd_title(text)
            
        wrapped_text = self._wrap_text_by_width(text, max_width=self.text_width, indent=self.text_indent)
        return wrapped_text + "\n\n"
    
    # 表格处理相关方法（保持原有复杂逻辑）
    def _get_merged_cell_info(self, cell, row_idx: int, col_idx: int) -> Tuple[int, int]:
        """获取单元格的合并信息"""
        vmerge = cell._tc.get_or_add_tcPr().first_child_found_in("w:vMerge")
        hmerge = cell._tc.get_or_add_tcPr().first_child_found_in("w:gridSpan")
        
        vspan = 1
        hspan = int(hmerge.val) if hmerge is not None else 1
        
        if vmerge is not None:
            current_text = cell.text.strip()
            
            if vmerge.val == "restart":
                current_row = row_idx + 1
                while current_row < len(self.current_table.rows):
                    next_cell = self.current_table.rows[current_row].cells[col_idx]
                    next_vmerge = next_cell._tc.get_or_add_tcPr().first_child_found_in("w:vMerge")
                    next_text = next_cell.text.strip()
                    
                    should_merge = (next_vmerge is not None and 
                                  (next_vmerge.val != "restart") and
                                  (not current_text or 
                                   not next_text or 
                                   current_text == next_text))
                    
                    if not should_merge:
                        break
                        
                    vspan += 1
                    current_row += 1
                    
            else:
                current_row = row_idx - 1
                found_start = False
                
                while current_row >= 0:
                    prev_cell = self.current_table.rows[current_row].cells[col_idx]
                    prev_vmerge = prev_cell._tc.get_or_add_tcPr().first_child_found_in("w:vMerge")
                    prev_text = prev_cell.text.strip()
                    
                    if (prev_vmerge is not None and 
                        prev_vmerge.val == "restart" and
                        (not current_text or 
                         not prev_text or 
                         current_text == prev_text)):
                        found_start = True
                        vspan = 0
                        break
                        
                    current_row -= 1
                    
                if not found_start:
                    vspan = 1
                
        return vspan, hspan
    
    def _process_cell_content(self, cell_text: str, total_width: int) -> List[str]:
        """处理单元格内容，进行文本换行"""
        if not cell_text:
            return ['']
        
        wrapped_lines = []
        available_width = total_width - CELL_PADDING
        
        lines = cell_text.split('\n')
        
        for line in lines:
            if not line:
                wrapped_lines.append('')
                continue
            
            current_line = []
            current_width = 0
            
            words = []
            temp_word = ""
            
            for char in line:
                if char.isspace():
                    if temp_word:
                        words.append(temp_word)
                        temp_word = ""
                    words.append(char)
                else:
                    if ord(char) > 127:
                        if temp_word:
                            words.append(temp_word)
                            temp_word = ""
                        words.append(char)
                    else:
                        temp_word += char
            if temp_word:
                words.append(temp_word)
            
            for word in words:
                word_width = self._get_string_width(word)
                
                if current_width + word_width <= available_width:
                    current_line.append(word)
                    current_width += word_width
                else:
                    if current_line:
                        wrapped_lines.append(''.join(current_line))
                    current_line = [word]
                    current_width = word_width
            
            if current_line:
                wrapped_lines.append(''.join(current_line))
        
        return wrapped_lines if wrapped_lines else ['']
    
    def _calculate_column_widths(self, table_data: List[List[Dict]], max_cols: int) -> List[int]:
        """计算表格每列的宽度"""
        col_widths = [0] * max_cols
        max_content_widths = [0] * max_cols
        
        for row_data in table_data:
            for col_idx, cell_data in enumerate(row_data):
                if cell_data['hspan'] == 1:
                    content_width = cell_data['max_line_width']
                    max_content_widths[col_idx] = max(max_content_widths[col_idx], content_width)
        
        for col_idx in range(max_cols):
            max_width = max_content_widths[col_idx]
            if max_width <= BASE_COLUMN_WIDTH:
                col_widths[col_idx] = BASE_COLUMN_WIDTH
            elif max_width <= BASE_COLUMN_WIDTH * LEVEL_2_MULTIPLIER:
                col_widths[col_idx] = BASE_COLUMN_WIDTH * LEVEL_2_MULTIPLIER
            elif max_width <= BASE_COLUMN_WIDTH * LEVEL_3_MULTIPLIER:
                col_widths[col_idx] = BASE_COLUMN_WIDTH * LEVEL_3_MULTIPLIER
            else:
                col_widths[col_idx] = BASE_COLUMN_WIDTH * LEVEL_3_MULTIPLIER
            
        for row_data in table_data:
            for col_idx, cell_data in enumerate(row_data):
                if cell_data['hspan'] > 1:
                    total_width = sum(col_widths[i] for i in range(col_idx, min(col_idx + cell_data['hspan'], len(col_widths))))
                    cell_data['available_width'] = total_width - CELL_PADDING
        
        return col_widths
    
    def _format_table_row(self, row_data: List[Dict], col_widths: List[int], line_idx: int) -> str:
        """格式化表格行"""
        row_content = '|'
        col_idx = 0
        
        while col_idx < len(row_data):
            cell_data = row_data[col_idx]
            hspan = cell_data['hspan']
            
            total_width = sum(col_widths[i] for i in range(col_idx, min(col_idx + hspan, len(col_widths)))) + (hspan - 1)
            
            if cell_data['vspan'] == 0:
                cell_text = ''
            else:
                cell_text = cell_data['lines'][line_idx] if line_idx < len(cell_data['lines']) else ''
            
            content_width = self._get_string_width(cell_text)
            left_padding = CELL_LEFT_PADDING
            right_padding = total_width - content_width - left_padding
            
            row_content += ' ' * left_padding + cell_text + ' ' * right_padding + '|'
            
            col_idx += hspan
        
        return row_content
    
    def _collect_table_data(self, table) -> Tuple[List[List[Dict]], int]:
        """收集表格数据和合并单元格信息"""
        max_cols = max(len(row.cells) for row in table.rows)
        table_data = []
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for col_idx, cell in enumerate(row.cells):
                vspan, hspan = self._get_merged_cell_info(cell, row_idx, col_idx)
                
                if cell._tc.get_or_add_tcPr().first_child_found_in("w:vMerge") is not None and \
                   cell._tc.get_or_add_tcPr().first_child_found_in("w:vMerge").val != "restart":
                    continue
                
                cell_text = self._clean_text(cell.text)
                original_lines = cell_text.split('\n')
                line_widths = [self._get_string_width(line) for line in original_lines]
                max_line_width = max(line_widths) if line_widths else 0
                
                needs_double_width = len(original_lines) - 1 >= NEWLINE_THRESHOLD
                
                row_data.append({
                    'text': cell_text,
                    'lines': original_lines,
                    'hspan': hspan,
                    'vspan': vspan,
                    'needs_double_width': needs_double_width,
                    'max_line_width': max_line_width
                })
            
            while len(row_data) < max_cols:
                row_data.append({
                    'text': '',
                    'lines': [''],
                    'hspan': 1,
                    'vspan': 1,
                    'needs_double_width': False,
                    'max_line_width': 0
                })
                
            table_data.append(row_data)
            
        return table_data, max_cols

    def _process_cell_wrapping(self, table_data: List[List[Dict]], col_widths: List[int]):
        """处理所有单元格的文本换行"""
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                if cell_data['vspan'] > 1:
                    h_total_width = sum(col_widths[i] for i in range(col_idx, min(col_idx + cell_data['hspan'], len(col_widths)))) + (cell_data['hspan'] - 1)
                    
                    if cell_data['text']:
                        cell_data['lines'] = self._process_cell_content(cell_data['text'], h_total_width)
                        
                    for v_idx in range(row_idx + 1, min(row_idx + cell_data['vspan'], len(table_data))):
                        table_data[v_idx][col_idx] = {
                            'text': '',
                            'lines': [''],
                            'hspan': cell_data['hspan'],
                            'vspan': 0,
                            'needs_double_width': False,
                            'max_line_width': 0,
                            'merged_from': (row_idx, col_idx)
                        }
        
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                if cell_data['vspan'] == 1:
                    if not cell_data['text']:
                        cell_data['lines'] = ['']
                        continue
                    
                    total_width = sum(col_widths[i] for i in range(col_idx, min(col_idx + cell_data['hspan'], len(col_widths)))) + (cell_data['hspan'] - 1)
                    cell_data['lines'] = self._process_cell_content(cell_data['text'], total_width)

    def _generate_table_string(self, table_data: List[List[Dict]], col_widths: List[int]) -> str:
        """生成最终的表格字符串"""
        result = []
        
        separator = '+'
        for width in col_widths:
            separator += '-' * width + '+'
        
        for row_idx, row_data in enumerate(table_data):
            max_lines = 1
            for cell_data in row_data:
                if cell_data['vspan'] != 0:
                    max_lines = max(max_lines, len(cell_data['lines']))
            
            for cell_data in row_data:
                if cell_data['vspan'] != 0:
                    if not cell_data['lines']:
                        cell_data['lines'] = ['']
                    while len(cell_data['lines']) < max_lines:
                        cell_data['lines'].append('')
            
            if (row_idx == 0 or 
                not any(cell_data['vspan'] == 0 for cell_data in row_data) or
                any(cell_data['vspan'] > 1 for cell_data in row_data)):
                result.append(separator)
            
            for line_idx in range(max_lines):
                result.append(self._format_table_row(row_data, col_widths, line_idx))
        
        result.append(separator)
        ascii_table = '\n'.join(result) + '\n'
        
        # 用代码块包裹 ASCII 表格，兼顾可读性与通用性
        return f"```text\n{ascii_table}```\n\n"

    def _process_table(self, table) -> str:
        """处理表格，保持原始格式和对齐方式"""
        try:
            if not table.rows:
                return "\n"
                
            self.current_table = table
            
            table_data, max_cols = self._collect_table_data(table)
            col_widths = self._calculate_column_widths(table_data, max_cols)
            self._process_cell_wrapping(table_data, col_widths)
            result = self._generate_table_string(table_data, col_widths)
            
            return result
            
        except Exception as e:
            self.logger.error(f"表格处理失败: {str(e)}")
            return "【表格处理失败】\n"
    
    def extract_content(self, docx_path: str) -> str:
        """提取DOCX内容并返回格式化文本"""
        file_name = docx_path.split('\\')[-1].split('/')[-1]
        self.logger.info(f"开始处理: {file_name}")
        
        try:
            doc = Document(docx_path)
            content = []
            
            paragraphs = list(doc.paragraphs)
            tables = list(doc.tables)
            p_index = 0
            t_index = 0
            
            total_elements = len(list(doc.element.body))
            processed_elements = 0
            table_count = 0
            
            for element in doc.element.body:
                if element.tag.endswith('p'):  # 段落
                    if p_index < len(paragraphs):
                        paragraph = paragraphs[p_index]
                        if paragraph.style and paragraph.style.name and 'heading' in paragraph.style.name.lower():
                            content.append(self._process_heading(paragraph))
                        else:
                            content.append(self._process_normal_paragraph(paragraph))
                        p_index += 1
                elif element.tag.endswith('tbl'):  # 表格
                    if t_index < len(tables):
                        content.append(self._process_table(tables[t_index]))
                        table_count += 1
                        t_index += 1
                
                processed_elements += 1
                # 只在25%、50%、75%、100%时显示进度
                progress = processed_elements / total_elements
                if progress >= 0.25 and not hasattr(self, '_progress_25'):
                    self.logger.info(f"处理进度: 25% (已处理{table_count}个表格)")
                    self._progress_25 = True
                elif progress >= 0.50 and not hasattr(self, '_progress_50'):
                    self.logger.info(f"处理进度: 50% (已处理{table_count}个表格)")
                    self._progress_50 = True
                elif progress >= 0.75 and not hasattr(self, '_progress_75'):
                    self.logger.info(f"处理进度: 75% (已处理{table_count}个表格)")
                    self._progress_75 = True
                elif processed_elements == total_elements:
                    self.logger.info(f"处理完成: 共{total_elements}个元素，{table_count}个表格")
                    # 清理进度标记
                    for attr in ['_progress_25', '_progress_50', '_progress_75']:
                        if hasattr(self, attr):
                            delattr(self, attr)
            
            return ''.join(content)
            
        except FileNotFoundError:
            self.logger.error(f"文件不存在: {file_name}")
            raise
        except KeyError as e:
            # 处理docx内部损坏的引用（如断开的书签链接）
            error_msg = str(e)
            if 'word/' in error_msg and 'bookmark' in error_msg.lower():
                self.logger.warning(f"文件包含损坏的书签引用，已跳过: {error_msg}")
            else:
                self.logger.warning(f"文件包含损坏的内部引用，已跳过: {error_msg}")
            # 返回已处理的内容，而非完全失败
            if content:
                self.logger.info(f"已提取部分内容: {len(content)}个元素")
                return ''.join(content)
            raise
        except Exception as e:
            self.logger.error(f"处理失败: {file_name} - {str(e)}")
            raise
